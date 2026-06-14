"""
Export service: generates PDF and DOCX files for LawBot documents.

Handles:
- Generated legal documents → PDF / DOCX
- Contract analysis reports → PDF
- Compliance reports → PDF / XLSX
"""
import io
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, Tuple

import structlog
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings

logger = structlog.get_logger(__name__)

EXPORTS_DIR = Path(settings.upload_dir) / "exports"


def _ensure_exports_dir() -> Path:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORTS_DIR


# ---------------------------------------------------------------------------
# DOCX helpers (python-docx)
# ---------------------------------------------------------------------------

def _build_docx_from_text(title: str, content: str) -> bytes:
    """Convert plain/markdown text to a basic DOCX document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph("")  # spacer

    # Parse content line by line
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect markdown-style headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf_from_text(title: str, content: str) -> bytes:
    """
    Convert plain text to PDF.
    Tries reportlab first, falls back to DOCX-as-PDF via LibreOffice if available,
    or returns a plain-text PDF via fpdf2.
    """
    try:
        return _pdf_via_fpdf2(title, content)
    except ImportError:
        pass

    try:
        return _pdf_via_reportlab(title, content)
    except ImportError:
        pass

    raise RuntimeError(
        "No PDF library available. Install fpdf2 or reportlab: "
        "pip install fpdf2  OR  pip install reportlab"
    )


def _pdf_via_fpdf2(title: str, content: str) -> bytes:
    """Generate PDF using fpdf2."""
    from fpdf import FPDF  # type: ignore

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(6)

    # Body
    pdf.set_font("Helvetica", size=10)
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            pdf.ln(4)
            continue

        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 7, stripped[4:])
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 8, stripped[3:])
            pdf.set_font("Helvetica", size=10)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 9, stripped[2:])
            pdf.set_font("Helvetica", size=10)
        else:
            pdf.multi_cell(0, 6, stripped)

    return pdf.output()


def _pdf_via_reportlab(title: str, content: str) -> bytes:
    """Generate PDF using reportlab."""
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib.units import cm  # type: ignore
    from reportlab.lib import colors  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
    from reportlab.lib.enums import TA_CENTER  # type: ignore

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=3 * cm, rightMargin=3 * cm,
                            topMargin=2.5 * cm, bottomMargin=2.5 * cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontSize=16,
        textColor=colors.HexColor("#1a1a2e"),
        alignment=TA_CENTER,
        spaceAfter=20,
    )
    body_style = styles["Normal"]
    h1_style = styles["Heading1"]
    h2_style = styles["Heading2"]
    h3_style = styles["Heading3"]

    story = [Paragraph(title, title_style), Spacer(1, 12)]

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], h3_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], h2_style))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], h1_style))
        else:
            story.append(Paragraph(stripped, body_style))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# EXPORT SERVICE
# ---------------------------------------------------------------------------

class ExportService:
    """
    Handles PDF and DOCX export for:
    - AI-generated legal documents
    - Contract analysis reports
    - Compliance reports
    """

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    # ------------------------------------------------------------------ #
    # GENERATED DOCUMENTS
    # ------------------------------------------------------------------ #

    async def export_generated_document(
        self,
        document_id: str,
        format: str = "docx",
    ) -> Tuple[Path, str, str]:
        """
        Export a GeneratedDocument to PDF or DOCX.

        Returns:
            (file_path, filename, relative_download_url)
        """
        from app.models.generated_document import GeneratedDocument

        result = await self.db.execute(
            select(GeneratedDocument).where(
                GeneratedDocument.id == uuid.UUID(document_id)
            )
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise ValueError(f"Generated document {document_id} not found")

        if not doc.content:
            raise ValueError(f"Document {document_id} has no generated content")

        title = doc.title or "Legal Document"
        content = doc.content

        file_bytes, ext = self._render(title, content, format)

        exports_dir = _ensure_exports_dir()
        safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)[:60]
        filename = f"{safe_title}_{document_id[:8]}.{ext}"
        file_path = exports_dir / filename

        file_path.write_bytes(file_bytes)

        # Update record with file path
        doc.file_path = str(file_path)
        self.db.add(doc)
        await self.db.commit()

        download_url = f"/api/v1/generator/documents/{document_id}/download/?format={format}"
        logger.info("Generated document exported", document_id=document_id, format=format, path=str(file_path))
        return file_path, filename, download_url

    # ------------------------------------------------------------------ #
    # CONTRACT ANALYSIS REPORTS
    # ------------------------------------------------------------------ #

    async def export_contract_analysis(
        self,
        analysis_id: str,
        format: str = "pdf",
    ) -> Tuple[Path, str]:
        """
        Build and export a contract analysis report.
        """
        content = await self._build_contract_report(analysis_id)
        title = f"Contract Analysis Report — {analysis_id[:8]}"

        file_bytes, ext = self._render(title, content, format)

        exports_dir = _ensure_exports_dir()
        filename = f"contract_analysis_{analysis_id[:8]}.{ext}"
        file_path = exports_dir / filename
        file_path.write_bytes(file_bytes)

        logger.info("Contract analysis exported", analysis_id=analysis_id, format=format)
        return file_path, filename

    async def _build_contract_report(self, analysis_id: str) -> str:
        """Build the text content for a contract analysis report."""
        from sqlalchemy import text

        # Fetch analysis data — using raw dict access to avoid tight model coupling
        try:
            from app.models.document import Document  # imported for relationship joins
        except ImportError:
            pass

        rows = await self.db.execute(
            text(
                "SELECT * FROM contract_analyses WHERE id = :id LIMIT 1"
            ),
            {"id": analysis_id},
        )
        row = rows.mappings().first()
        if not row:
            raise ValueError(f"Contract analysis {analysis_id} not found")

        lines = [
            f"# Contract Analysis Report",
            f"",
            f"**Analysis ID:** {analysis_id}",
            f"**Generated:** {datetime.now(timezone.utc).strftime('%d %B %Y, %I:%M %p IST')}",
            f"**Overall Risk Level:** {row.get('overall_risk_level', 'N/A')}",
            f"**Risk Score:** {row.get('risk_score', 'N/A')}",
            f"**Jurisdiction:** {row.get('jurisdiction', 'India')}",
            f"",
            f"## Executive Summary",
            f"",
            str(row.get("executive_summary", "No summary available.")),
            f"",
        ]

        # Risk flags
        flag_rows = await self.db.execute(
            text("SELECT * FROM risk_flags WHERE analysis_id = :id ORDER BY severity"),
            {"id": analysis_id},
        )
        flags = flag_rows.mappings().all()
        if flags:
            lines += [f"## Risk Flags ({len(flags)} total)", ""]
            for f in flags:
                lines += [
                    f"### [{f.get('severity', '').upper()}] {f.get('title', '')}",
                    f"**Category:** {f.get('category', '')}",
                    f"**Description:** {f.get('description', '')}",
                    f"**Recommendation:** {f.get('recommendation', '')}",
                    "",
                ]

        # Clause analyses
        clause_rows = await self.db.execute(
            text("SELECT * FROM clause_analyses WHERE analysis_id = :id ORDER BY clause_name"),
            {"id": analysis_id},
        )
        clauses = clause_rows.mappings().all()
        if clauses:
            lines += [f"## Clause Analysis ({len(clauses)} clauses)", ""]
            for c in clauses:
                lines += [
                    f"### {c.get('clause_name', '')}",
                    f"**Risk Level:** {c.get('risk_level', '')}",
                    f"**Analysis:** {c.get('analysis', '')}",
                    "",
                ]

        # Recommendations
        recommendations = row.get("recommendations")
        if recommendations:
            lines += ["## Recommendations", ""]
            if isinstance(recommendations, list):
                for i, rec in enumerate(recommendations, 1):
                    lines.append(f"{i}. {rec}")
            else:
                lines.append(str(recommendations))
            lines.append("")

        lines += [
            "---",
            "*This report is generated by LawBot AI. It is not a substitute for professional legal advice.*",
            "*Please consult a qualified legal professional before acting on this analysis.*",
        ]

        return "\n".join(lines)

    # ------------------------------------------------------------------ #
    # COMPLIANCE REPORTS
    # ------------------------------------------------------------------ #

    async def export_compliance_report(
        self,
        user_id: str,
        framework: Optional[str],
        format: str = "pdf",
        include_recommendations: bool = True,
    ) -> Tuple[Path, str]:
        """Build and export a compliance report for a user/organization."""
        content = await self._build_compliance_report(
            user_id, framework, include_recommendations
        )
        framework_str = f"_{framework}" if framework else ""
        title = f"Compliance Report{framework_str.replace('_', ' ')}"

        file_bytes, ext = self._render(title, content, format)

        exports_dir = _ensure_exports_dir()
        filename = f"compliance_report_{user_id[:8]}{framework_str}_{datetime.now(timezone.utc).strftime('%Y%m%d')}.{ext}"
        file_path = exports_dir / filename
        file_path.write_bytes(file_bytes)

        logger.info("Compliance report exported", user_id=user_id, format=format)
        return file_path, filename

    async def _build_compliance_report(
        self,
        user_id: str,
        framework: Optional[str],
        include_recommendations: bool,
    ) -> str:
        """Build the text content for a compliance report."""
        from sqlalchemy import text

        filter_sql = "WHERE user_id = :user_id"
        params: dict = {"user_id": user_id}
        if framework:
            filter_sql += " AND framework = :framework"
            params["framework"] = framework

        item_rows = await self.db.execute(
            text(f"SELECT * FROM compliance_items {filter_sql} ORDER BY priority, due_date"),
            params,
        )
        items = item_rows.mappings().all()

        lines = [
            "# Compliance Report",
            "",
            f"**Generated:** {datetime.now(timezone.utc).strftime('%d %B %Y')}",
            f"**Framework:** {framework or 'All Frameworks'}",
            f"**Total Items:** {len(items)}",
            "",
        ]

        # Group by framework
        by_framework: dict = {}
        for item in items:
            fw = item.get("framework", "Unknown")
            by_framework.setdefault(fw, []).append(item)

        for fw, fw_items in by_framework.items():
            completed = sum(1 for i in fw_items if i.get("status") == "completed")
            lines += [
                f"## {fw.upper()} — {completed}/{len(fw_items)} completed",
                "",
            ]
            for item in fw_items:
                status_icon = "✓" if item.get("status") == "completed" else "○"
                lines += [
                    f"### {status_icon} {item.get('title', '')}",
                    f"**Status:** {item.get('status', '')}  |  **Priority:** {item.get('priority', '')}",
                    f"**Due Date:** {item.get('due_date', 'N/A')}",
                ]
                if item.get("description"):
                    lines.append(f"**Description:** {item['description']}")
                if include_recommendations and item.get("notes"):
                    lines.append(f"**Notes:** {item['notes']}")
                lines.append("")

        lines += [
            "---",
            "*This report is generated by LawBot AI. Consult a qualified CA/CS/lawyer for professional advice.*",
        ]

        return "\n".join(lines)

    # ------------------------------------------------------------------ #
    # INTERNAL RENDER DISPATCH
    # ------------------------------------------------------------------ #

    def _render(self, title: str, content: str, format: str) -> Tuple[bytes, str]:
        """Dispatch to the correct renderer based on format."""
        fmt = format.lower().strip(".")
        if fmt == "docx":
            return _build_docx_from_text(title, content), "docx"
        elif fmt == "pdf":
            return _build_pdf_from_text(title, content), "pdf"
        else:
            raise ValueError(f"Unsupported export format: {format}. Use 'pdf' or 'docx'.")

    # ------------------------------------------------------------------ #
    # SERVE EXPORT FILE (used by route handlers)
    # ------------------------------------------------------------------ #

    async def get_export_file(self, filename: str) -> Tuple[Path, str]:
        """
        Retrieve a previously-exported file by filename.
        Returns (file_path, media_type).
        """
        file_path = EXPORTS_DIR / filename
        if not file_path.exists():
            raise FileNotFoundError(f"Export file not found: {filename}")

        ext = file_path.suffix.lower()
        media_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }
        media_type = media_types.get(ext, "application/octet-stream")
        return file_path, media_type
